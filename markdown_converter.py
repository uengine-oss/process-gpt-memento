#!/usr/bin/env python3
"""
Unified Markdown Converter
마크다운 텍스트를 HTML 또는 DOCX로 변환하는 통합 함수
"""

import re
import markdown
import subprocess
import tempfile
import os
from typing import Optional, Union

# Optional imports for DOCX support
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def check_mermaid_cli() -> bool:
    """mermaid-cli가 설치되어 있는지 확인"""
    try:
        result = subprocess.run(['mmdc', '--version'], 
                              capture_output=True, text=True)
        return result.returncode == 0
    except FileNotFoundError:
        return False

def generate_mermaid_svg(mermaid_code: str) -> str:
    """Mermaid 코드를 SVG로 변환"""
    if not check_mermaid_cli():
        return f'<div class="error">Mermaid CLI not available. Install with: npm install -g @mermaid-js/mermaid-cli</div>'
    
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as temp_mmd:
            temp_mmd.write(mermaid_code)
            temp_mmd_path = temp_mmd.name
        
        temp_svg_path = temp_mmd_path.replace('.mmd', '.svg')
        
        result = subprocess.run([
            'mmdc', '-i', temp_mmd_path, '-o', temp_svg_path, 
            '--backgroundColor', 'transparent'
        ], capture_output=True, text=True)
        
        if result.returncode == 0 and os.path.exists(temp_svg_path):
            with open(temp_svg_path, 'r', encoding='utf-8') as svg_file:
                svg_content = svg_file.read()
            
            os.unlink(temp_mmd_path)
            os.unlink(temp_svg_path)
            return svg_content
        else:
            os.unlink(temp_mmd_path)
            return f'<div class="error">Failed to generate diagram: {result.stderr}</div>'
            
    except Exception as e:
        return f'<div class="error">Error generating diagram: {e}</div>'

def convert_markdown_to_html(markdown_text: str, mode: str = 'inline', 
                           title: str = "Markdown Document") -> str:
    """
    마크다운 텍스트를 HTML로 변환
    
    Args:
        markdown_text: 변환할 마크다운 텍스트
        mode: 변환 모드 ('inline', 'dynamic')
        title: HTML 문서 제목
    
    Returns:
        HTML 문자열
    """
    
    if mode == 'inline':
        return _convert_to_html_inline(markdown_text, title)
    elif mode == 'dynamic':
        return _convert_to_html_dynamic(markdown_text, title)
    else:
        raise ValueError(f"Unknown mode: {mode}. Use 'inline' or 'dynamic'")

def _convert_to_html_inline(markdown_text: str, title: str) -> str:
    """인라인 SVG 방식으로 HTML 변환"""
    
    def replace_mermaid_with_svg(match):
        mermaid_code = match.group(1).strip()
        svg_content = generate_mermaid_svg(mermaid_code)
        return f'<div class="mermaid-container">\n{svg_content}\n</div>'
    
    # Mermaid 블록을 SVG로 대체
    processed_content = re.sub(r'```mermaid\n(.*?)\n```', 
                             replace_mermaid_with_svg, 
                             markdown_text, flags=re.DOTALL)
    
    # 마크다운을 HTML로 변환
    md = markdown.Markdown(extensions=[
        'extra', 'codehilite', 'toc', 'fenced_code', 'tables'
    ])
    
    html_content = md.convert(processed_content)
    
    return _create_html_template(html_content, title, include_mermaid_js=False)

def _convert_to_html_dynamic(markdown_text: str, title: str) -> str:
    """동적 Mermaid.js 방식으로 HTML 변환"""
    
    def replace_mermaid_blocks(match):
        mermaid_code = match.group(1)
        return f'<div class="mermaid">\n{mermaid_code}\n</div>'
    
    # Mermaid 블록을 div로 대체
    processed_content = re.sub(r'```mermaid\n(.*?)\n```', 
                             replace_mermaid_blocks, 
                             markdown_text, flags=re.DOTALL)
    
    # 마크다운을 HTML로 변환
    md = markdown.Markdown(extensions=[
        'extra', 'codehilite', 'toc', 'fenced_code', 'tables'
    ])
    
    html_content = md.convert(processed_content)
    
    return _create_html_template(html_content, title, include_mermaid_js=True)

def _create_html_template(html_content: str, title: str, 
                        include_mermaid_js: bool = False) -> str:
    """HTML 템플릿 생성"""
    
    mermaid_script = ""
    if include_mermaid_js:
        mermaid_script = """
    <!-- Mermaid for diagrams -->
    <script type="module">
        import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs';
        mermaid.initialize({ 
            startOnLoad: true,
            theme: 'default',
            securityLevel: 'loose',
            flowchart: {
                useMaxWidth: true,
                htmlLabels: true
            }
        });
    </script>"""
    
    return f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    
    <!-- MathJax for mathematical expressions -->
    <script src="https://polyfill.io/v3/polyfill.min.js?features=es6"></script>
    <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
    <script>
        window.MathJax = {{
            tex: {{
                inlineMath: [['$', '$'], ['\\\\(', '\\\\)']],
                displayMath: [['$$', '$$'], ['\\\\[', '\\\\]']]
            }}
        }};
    </script>
    
    {mermaid_script}
    
    <!-- Prism.js for syntax highlighting -->
    <link href="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism.min.css" rel="stylesheet" />
    <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/prism.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-python.min.js"></script>
    
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
        }}
        
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 2rem;
            margin-bottom: 1rem;
        }}
        
        h1 {{
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        
        h2 {{
            border-bottom: 2px solid #e74c3c;
            padding-bottom: 8px;
        }}
        
        code {{
            background-color: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Monaco', 'Menlo', 'Ubuntu Mono', monospace;
        }}
        
        pre {{
            background-color: #f8f8f8;
            border: 1px solid #e1e1e8;
            border-radius: 5px;
            padding: 1rem;
            overflow-x: auto;
        }}
        
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1rem 0;
        }}
        
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        
        th {{
            background-color: #f2f2f2;
            font-weight: bold;
        }}
        
        .mermaid, .mermaid-container {{
            text-align: center;
            margin: 2rem 0;
            padding: 1rem;
            background-color: #fafafa;
            border: 1px solid #e1e1e8;
            border-radius: 8px;
        }}
        
        .mermaid-container svg {{
            max-width: 100%;
            height: auto;
        }}
        
        .error {{
            color: #e74c3c;
            background-color: #fdf2f2;
            border: 1px solid #f5c6cb;
            border-radius: 4px;
            padding: 1rem;
            margin: 1rem 0;
        }}
        
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1rem auto;
            border: 1px solid #ddd;
            border-radius: 5px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        
        blockquote {{
            border-left: 4px solid #3498db;
            margin: 1rem 0;
            padding: 0.5rem 1rem;
            background-color: #f8f9fa;
        }}
        
        .toc {{
            background-color: #f8f9fa;
            border: 1px solid #e9ecef;
            border-radius: 5px;
            padding: 1rem;
            margin: 1rem 0;
        }}
        
        hr {{
            border: none;
            height: 2px;
            background: linear-gradient(to right, #3498db, #e74c3c);
            margin: 2rem 0;
        }}
        
        @media (max-width: 768px) {{
            body {{
                padding: 10px;
            }}
            
            h1 {{
                font-size: 1.8rem;
            }}
            
            h2 {{
                font-size: 1.5rem;
            }}
            
            h3 {{
                font-size: 1.3rem;
            }}
        }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""

def convert_markdown_to_docx(markdown_text: str, file_name: Optional[str] = None,
                            output_path: Optional[str] = None,
                            image_format: str = 'png') -> Union[str, bytes]:
    """
    마크다운 텍스트를 DOCX로 변환 (한글 지원 개선)
    
    Args:
        markdown_text: 변환할 마크다운 텍스트
        file_name: 출력 파일 이름 (None이면 루트 경로에 저장하고 바이트 반환)
        output_path: 출력 파일 경로 (None이면 루트 경로에 저장하고 바이트 반환)
        image_format: 이미지 형식 ('png', 'svg', 'pdf')
    
    Returns:
        파일 경로 또는 바이트 데이터
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX conversion. Install with: pip install python-docx")
    
    # 한글 폰트 설정을 위한 Document 생성
    doc = Document()
    
    # 한글 폰트 설정
    _set_korean_font(doc)
    
    # 임시 디렉토리 생성
    with tempfile.TemporaryDirectory() as temp_dir:
        # Mermaid 다이어그램 처리
        image_files = []
        image_counter = 1
        
        def replace_mermaid_block(match):
            nonlocal image_counter
            mermaid_code = match.group(1).strip()
            
            image_filename = f"mermaid_diagram_{image_counter}.{image_format}"
            image_path = os.path.join(temp_dir, image_filename)
            
            if _generate_mermaid_image(mermaid_code, image_path, image_format):
                image_files.append(image_path)
                image_counter += 1
                return f"MERMAID_IMAGE:{image_path}"
            else:
                return f"**[Mermaid Diagram Error: Could not generate image]**"
        
        # Mermaid 블록을 플레이스홀더로 대체
        processed_content = re.sub(r'```mermaid\n(.*?)\n```', 
                                 replace_mermaid_block, 
                                 markdown_text, flags=re.DOTALL)
        
        # 마크다운 파싱 및 DOCX 생성
        _parse_markdown_to_docx(doc, processed_content)
        
        # 문서 저장
        if output_path:
            if file_name:
                output_path = os.path.join(output_path, file_name)
            doc.save(output_path)
            return output_path
        else:
            # 루트 경로에 기본 파일명으로 저장
            from datetime import datetime
            if file_name:
                default_filename = file_name
            else:
                default_filename = f"converted_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            root_path = os.path.join(os.getcwd(), default_filename)
            doc.save(root_path)
            
            # 바이트로도 반환
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

def _set_korean_font(doc):
    """한글 폰트 설정"""
    try:
        # 문서의 기본 스타일 설정
        style = doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'  # Windows 한글 폰트
        font.size = Pt(11)
        
        # 제목 스타일들도 한글 폰트로 설정
        for i in range(1, 10):
            try:
                heading_style = doc.styles[f'Heading {i}']
                heading_font = heading_style.font
                heading_font.name = '맑은 고딕'
                heading_font.size = Pt(16 - i)  # 제목 크기 조정
            except:
                continue
                
    except Exception as e:
        print(f"폰트 설정 중 오류: {e}")

def _generate_mermaid_image(mermaid_code: str, output_path: str, 
                          image_format: str) -> bool:
    """Mermaid 이미지 생성 (한글 지원 개선)"""
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as temp_mmd:
            temp_mmd.write(mermaid_code)
            temp_mmd_path = temp_mmd.name
        
        result = subprocess.run([
            'mmdc', '-i', temp_mmd_path, '-o', output_path,
            '--backgroundColor', 'white',
            '--width', '1200',
            '--height', '800'
        ], capture_output=True, text=True, encoding='utf-8')
        
        os.unlink(temp_mmd_path)
        return result.returncode == 0 and os.path.exists(output_path)
        
    except Exception as e:
        print(f"Exception in Mermaid image generation: {e}")
        return False

def _parse_markdown_to_docx(doc, content: str):
    """마크다운을 DOCX로 파싱 (한글 지원 개선)"""
    lines = content.split('\n')
    
    current_paragraph = None
    in_code_block = False
    in_list = False
    list_level = 0
    
    for line in lines:
        original_line = line
        line = line.strip()
        
        if not line and not in_code_block:
            current_paragraph = None
            in_list = False
            continue
        
        # 헤딩 처리
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            heading = doc.add_heading(heading_text, level)
            
            # 제목에 한글 폰트 적용
            for run in heading.runs:
                run.font.name = '맑은 고딕'
                run.font.size = Pt(16 - level)
            
            current_paragraph = None
            in_list = False
        
        # 리스트 처리
        elif line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            paragraph = doc.add_paragraph()
            paragraph.style = 'List Bullet'
            run = paragraph.add_run(list_text)
            run.font.name = '맑은 고딕'
            current_paragraph = paragraph
            in_list = True
        
        # 번호 리스트 처리
        elif re.match(r'^\d+\.\s', line):
            list_text = re.sub(r'^\d+\.\s', '', line).strip()
            paragraph = doc.add_paragraph()
            paragraph.style = 'List Number'
            run = paragraph.add_run(list_text)
            run.font.name = '맑은 고딕'
            current_paragraph = paragraph
            in_list = True
        
        # 코드 블록 처리
        elif line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                paragraph = doc.add_paragraph()
                paragraph.style = 'No Spacing'
                run = paragraph.add_run('코드 블록:')
                run.font.name = '맑은 고딕'
                run.bold = True
            current_paragraph = None
        
        # Mermaid 이미지 플레이스홀더 처리
        elif line.startswith('MERMAID_IMAGE:'):
            image_path = line.replace('MERMAID_IMAGE:', '')
            if os.path.exists(image_path):
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()
                try:
                    run.add_picture(image_path, width=Inches(6))
                except Exception as e:
                    error_run = paragraph.add_run(f"[이미지를 삽입할 수 없습니다: {os.path.basename(image_path)}]")
                    error_run.font.name = '맑은 고딕'
            current_paragraph = None
        
        # 테이블 처리
        elif line.startswith('|') and line.endswith('|'):
            # 테이블 처리 로직 (간단한 구현)
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if not hasattr(doc, '_current_table'):
                doc._current_table = doc.add_table(rows=1, cols=len(cells))
                doc._current_table.style = 'Table Grid'
                header_cells = doc._current_table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    header_cells[i].text = cell_text
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '맑은 고딕'
                            run.bold = True
            else:
                row_cells = doc._current_table.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell_text
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '맑은 고딕'
            current_paragraph = None
        
        # 일반 텍스트 처리
        elif not in_code_block:
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
            
            # 볼드 텍스트 처리
            if '**' in original_line:
                parts = original_line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        run = current_paragraph.add_run(part)
                        run.font.name = '맑은 고딕'
                    else:
                        run = current_paragraph.add_run(part)
                        run.font.name = '맑은 고딕'
                        run.bold = True
            else:
                run = current_paragraph.add_run(original_line + ' ')
                run.font.name = '맑은 고딕'

# 사용 예시
if __name__ == "__main__":
    # 예시 마크다운 텍스트
    sample_markdown = """# 제목

이것은 **볼드 텍스트**입니다.

## 서브제목

```mermaid
graph TD
    A[시작] --> B[처리]
    B --> C[종료]
```

### 코드 예시

```python
print("Hello, World!")
```

- 리스트 항목 1
- 리스트 항목 2

| 컬럼1 | 컬럼2 |
|-------|-------|
| 데이터1 | 데이터2 |
"""
    
    # HTML 변환 예시
    html_result = convert_markdown_to_html(sample_markdown, mode='inline', title="샘플 문서")
    print("HTML 변환 완료!")
    
    # HTML 파일로 저장
    with open("sample_output.html", "w", encoding="utf-8") as f:
        f.write(html_result)
    print("HTML 파일 저장 완료: sample_output.html")
    
    # DOCX 변환 예시 (파일로 저장)
    try:
        docx_path = convert_markdown_to_docx(sample_markdown, "sample_output.docx")
        print(f"DOCX 변환 완료: {docx_path}")
    except ImportError as e:
        print(f"DOCX 변환 실패: {e}")
    
    # DOCX 변환 예시 (바이트로 반환)
    try:
        docx_bytes = convert_markdown_to_docx(sample_markdown)
        print(f"DOCX 바이트 크기: {len(docx_bytes)} bytes")
    except ImportError as e:
        print(f"DOCX 바이트 변환 실패: {e}")
