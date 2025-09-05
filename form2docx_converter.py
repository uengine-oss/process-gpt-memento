from bs4 import BeautifulSoup
from docx import Document

def form_to_docx(html: str, values: dict) -> bytes:
    """
    form html + form value(JSON)을 Word 문서(docx)로 변환

    Args:
        html: form html (string)
        values: form value (dict)
    Returns:
        bytes: docx 파일의 bytes
    """
    doc = Document()
    soup = BeautifulSoup(html, "html.parser")

    # 루트 키 추출 (예: contract_management_process_request_activity_form)
    root_key = next(iter(values.keys()))
    form_values = values[root_key]

    for section in soup.find_all("section"):
        row_layout = section.find("row-layout")
        if not row_layout:
            continue

        # 섹션 제목
        section_title = row_layout.get("alias", "")
        if section_title:
            doc.add_heading(section_title, level=2)

        # row 단위 → Word 테이블
        for row in section.find_all("div", class_="row"):
            cols = row.find_all("div", class_=lambda x: x and x.startswith("col-sm-"))
            if not cols:
                continue

            # Word Table: 라벨 행 + 값 행
            table = doc.add_table(rows=2, cols=len(cols))
            table.autofit = True

            label_cells = table.rows[0].cells
            value_cells = table.rows[1].cells

            for i, col in enumerate(cols):
                field = col.find()
                if not field:
                    continue
                alias = field.get("alias", "")
                key = field.get("name")
                value = extract_value(field, key, form_values)

                label_cells[i].text = alias
                value_cells[i].text = value

    # docx → bytes 변환
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def extract_value(field, key, values):
    """필드 타입별 값 추출"""
    v = values.get(key)
    if v is None:
        return ""

    ftype = field.name

    if ftype == "select-field":
        try:
            items = eval(field.get("items", "[]"))
            for item in items:
                for k, v2 in item.items():
                    if k == v:
                        return v2
        except Exception:
            pass
        return str(v)

    elif ftype == "textarea-field":
        return str(v)

    elif ftype == "file-field":
        if isinstance(v, dict):
            return v.get("name", "")
        elif isinstance(v, list):
            return ", ".join([f.get("name", "") for f in v if isinstance(f, dict)])
        return str(v)

    elif ftype == "user-select-field":
        if isinstance(v, list):
            return ", ".join([u.get("username", "") for u in v])
        return str(v)

    else:
        return str(v)
