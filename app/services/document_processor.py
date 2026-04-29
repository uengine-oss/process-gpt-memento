"""
Document loader and processor
"""
import os
import re
import docx
import uuid
import tempfile
import asyncio
import io
import zipfile
from typing import List, Optional, Dict, Any, Tuple, AsyncIterator
from pathlib import Path
from pydantic import BaseModel
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from openai import OpenAI

from app.plugins.chunkers import get_chunker
from app.plugins.parsers import (
    get_pdf_parser,
    get_synap_parser,
    synap_supports,
)
from langchain_community.document_loaders import (
    UnstructuredWordDocumentLoader,
    UnstructuredPowerPointLoader,
    UnstructuredExcelLoader,
    UnstructuredFileLoader,
    PyPDFLoader,
    TextLoader
)
import fitz  # PyMuPDF for PDF image extraction

# Allow loading vendored extract_hwp when the installed extract-hwp package has no module (PyPI 0.1.0 packaging bug)
_vendor_dir = Path(__file__).resolve().parent / "vendor"
if _vendor_dir.is_dir() and str(_vendor_dir) not in __import__("sys").path:
    __import__("sys").path.insert(0, str(_vendor_dir))


def _extract_text_from_hwp_or_hwpx(file_path: str, file_extension: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text from HWP or HWPX file. Uses extract_hwp if available; otherwise .hwpx via zip+xml.
    PyPI extract-hwp 0.1.0 ships no module (packaging bug). For .hwp either install from Git or add
    vendor: clone https://github.com/thlee/extract-hwp and copy src/extract_hwp into vendor/.
    Returns (text, error_message); error_message is None on success.
    """
    try:
        from extract_hwp import extract_text_from_hwp
        return extract_text_from_hwp(file_path)
    except ModuleNotFoundError:
        pass  # extract-hwp not installed or broken (PyPI 0.1.0 has no module); use fallback

    if file_extension == ".hwp":
        return (
            None,
            "HWP 파일 처리를 위해 extract-hwp를 GitHub에서 설치해 주세요: uv pip install \"extract-hwp @ git+https://github.com/thlee/extract-hwp.git\" (또는 vendor 폴더에 소스 추가)",
        )

    if file_extension == ".hwpx":
        try:
            import xml.etree.ElementTree as ET

            def _ltag(elem):
                t = elem.tag
                return t.split('}', 1)[1] if '}' in t else t

            def _collect_t(elem):
                parts = []
                for n in elem.iter():
                    if _ltag(n) == 'tbl':
                        continue
                    if _ltag(n) == 't' and n.text:
                        parts.append(n.text)
                return "".join(parts)

            def _tbl_to_md(tbl):
                cells = []
                for tr in tbl:
                    if _ltag(tr) != 'tr':
                        continue
                    for tc in tr:
                        if _ltag(tc) != 'tc':
                            continue
                        row = col = 0
                        for cc in tc:
                            tag = _ltag(cc)
                            if tag == 'cellAddr':
                                for k, v in cc.attrib.items():
                                    if 'colAddr' in k: col = int(v)
                                    if 'rowAddr' in k: row = int(v)
                        tp = []
                        for sub in tc.iter():
                            if _ltag(sub) == 't' and sub.text:
                                tp.append(sub.text)
                        cells.append((row, col, " ".join("".join(tp).split())))
                if not cells:
                    return ""
                mr = max(r + 1 for r, c, t in cells)
                mc = max(c + 1 for r, c, t in cells)
                grid = [[""] * mc for _ in range(mr)]
                for r, c, t in cells:
                    grid[r][c] = t
                cw = [max(3, *(len(grid[r][c]) for r in range(mr))) for c in range(mc)]
                lines = []
                for r in range(mr):
                    lines.append("| " + " | ".join(grid[r][c].ljust(cw[c]) for c in range(mc)) + " |")
                    if r == 0:
                        lines.append("| " + " | ".join("-" * cw[c] for c in range(mc)) + " |")
                return "\n".join(lines)

            def _walk(elem, results):
                tag = _ltag(elem)
                if tag == 'tbl':
                    md = _tbl_to_md(elem)
                    if md:
                        results.append(md)
                    return
                if tag == 'p':
                    has_tbl = any(_ltag(d) == 'tbl' for d in elem.iter() if d is not elem)
                    if has_tbl:
                        for ch in elem:
                            _walk(ch, results)
                    else:
                        text = _collect_t(elem)
                        if text.strip():
                            results.append(text)
                    return
                for ch in elem:
                    _walk(ch, results)

            with zipfile.ZipFile(file_path, "r") as z:
                names = z.namelist()
                section_files = sorted(n for n in names if "Contents/section" in n and n.endswith(".xml"))
                if not section_files:
                    content_name = next((n for n in names if "contents" in n.lower() and n.endswith(".xml")), None)
                    if not content_name:
                        return (None, "HWPX: section or contents XML not found")
                    with z.open(content_name) as f:
                        raw = f.read().decode("utf-8", errors="replace")
                    text = re.sub(r"<[^>]+>", " ", raw)
                    text = re.sub(r"\s+", " ", text).strip()
                    return (text, None)
                sections = []
                for section_file in section_files:
                    with z.open(section_file) as f:
                        root = ET.fromstring(f.read())
                    parts = []
                    _walk(root, parts)
                    if parts:
                        sections.append("\n\n".join(parts))
                return ("\n\n".join(sections), None)
        except Exception as e:
            return (None, str(e))

    return (None, f"Unsupported extension: {file_extension}")


class DocumentProcessor:
    def __init__(self, chunk_size: int = 2000, chunk_overlap: int = 400):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        # Chunking strategy is selected via CHUNKER_STRATEGY env var
        # (recursive / fixed_token / markdown_header / semantic / hybrid).
        self.chunker = get_chunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        self._openai_client: Optional[OpenAI] = None

    def _get_openai_client(self) -> OpenAI:
        if self._openai_client is None:
            self._openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        return self._openai_client

    async def _generate_section_title_single(
        self, content: str, semaphore: asyncio.Semaphore
    ) -> str:
        """청크 1개의 section_title을 Structured Output으로 생성한다."""

        class _TitleResponse(BaseModel):
            title: str

        snippet = content[:300].replace("\n", " ")
        async with semaphore:
            try:
                client = self._get_openai_client()
                response = await asyncio.to_thread(
                    client.beta.chat.completions.parse,
                    model="gpt-4o-mini",
                    messages=[{
                        "role": "user",
                        "content": (
                            "다음 문서 내용에 어울리는 소제목(10자 이내)을 생성하세요.\n\n"
                            f"{snippet}"
                        ),
                    }],
                    response_format=_TitleResponse,
                    temperature=0,
                    max_tokens=50,
                )
                parsed = response.choices[0].message.parsed
                return parsed.title.strip() if parsed else ""
            except Exception as e:
                print(f"section_title 생성 실패: {e}")
                return ""

    async def _generate_section_titles(self, chunks: List[Document]) -> List[str]:
        """청크별 section_title을 병렬로 생성한다 (Structured Output 사용).

        동시 요청 수를 semaphore로 제한해 rate limit를 방지한다.
        """
        semaphore = asyncio.Semaphore(10)
        tasks = [
            self._generate_section_title_single(chunk.page_content or "", semaphore)
            for chunk in chunks
        ]
        return list(await asyncio.gather(*tasks))

    def _load_docx_with_python_docx(self, tmp_path: str, file_name: str) -> List[Document]:
        """DOCX 텍스트 추출 paragraph + table 셀 텍스트."""
        if docx is None:
            raise RuntimeError("python-docx is not installed")
        doc = docx.Document(tmp_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = "\n\n".join(parts) if parts else ""
        return [Document(page_content=text)]

    async def load_document(self, file_content: bytes, file_name: str) -> Optional[List[Document]]:
        """Async: Load a document from memory (BytesIO object)."""
        try:
            file_extension = os.path.splitext(file_name)[1].lower()

            documents = None

            # Synap 원격 파서가 활성화되어 있고 지원 확장자이면 우선 시도.
            # 실패 시 아래 로컬 파서 분기로 자동 폴백한다.
            if synap_supports(file_extension):
                data = await asyncio.to_thread(file_content.read)
                try:
                    documents = await get_synap_parser().parse(data, file_name)
                    print(f"[synap] '{file_name}' 원격 파싱 성공 (pages={len(documents)})")
                except Exception as e:
                    print(f"[synap] '{file_name}' 원격 파싱 실패 → 로컬 파서로 폴백: {e}")
                    documents = None
                    try:
                        file_content.seek(0)
                    except Exception:
                        pass

            if documents is not None:
                pass
            elif file_extension == '.txt':
                content = await asyncio.to_thread(file_content.read)
                content = content.decode('utf-8-sig')
                documents = [Document(page_content=content)]
            elif file_extension == '.docx':
                # Save BytesIO to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    await asyncio.to_thread(tmp.write, file_content.read())
                    tmp_path = tmp.name
                try:
                    documents = await asyncio.to_thread(
                        self._load_docx_with_python_docx,
                        tmp_path,
                        file_name
                    )
                finally:
                    await asyncio.to_thread(os.unlink, tmp_path)
            elif file_extension == '.pptx':
                # Save BytesIO to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pptx') as tmp:
                    await asyncio.to_thread(tmp.write, file_content.read())
                    tmp_path = tmp.name
                try:
                    loader = UnstructuredPowerPointLoader(tmp_path, mode="single")
                    documents = await asyncio.to_thread(loader.load)
                finally:
                    await asyncio.to_thread(os.unlink, tmp_path)
            elif file_extension == '.xlsx':
                # Save BytesIO to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
                    await asyncio.to_thread(tmp.write, file_content.read())
                    tmp_path = tmp.name
                try:
                    loader = UnstructuredExcelLoader(tmp_path, mode="single")
                    documents = await asyncio.to_thread(loader.load)
                finally:
                    await asyncio.to_thread(os.unlink, tmp_path)
            elif file_extension == '.pdf':
                data = await asyncio.to_thread(file_content.read)
                documents = await get_pdf_parser().parse(data, file_name)
            elif file_extension in ('.hwp', '.hwpx'):
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp:
                    await asyncio.to_thread(tmp.write, file_content.read())
                    tmp_path = tmp.name
                try:
                    text, error = await asyncio.to_thread(
                        _extract_text_from_hwp_or_hwpx, tmp_path, file_extension
                    )
                    if error is not None:
                        print(f"HWP/HWPX extraction error for {file_name}: {error}")
                        return None
                    documents = [Document(page_content=text or "")]
                finally:
                    await asyncio.to_thread(os.unlink, tmp_path)
            else:
                print(f"Unsupported file type: {file_extension}")
                return None

            # Add basic metadata including UUID
            doc_id = str(uuid.uuid4())  # Generate a single UUID for the document
            for doc in documents:
                doc.metadata.update({
                    "id": doc_id,
                    "source": file_name,
                    "file_name": file_name,
                    "file_type": file_extension[1:],
                    "language": "ko",
                    "content_length": len(doc.page_content)
                })
            
            return documents
        except Exception as e:
            print(f"Error loading document from memory {file_name}: {e}")
            return None

    @staticmethod
    def _attach_chunk_bbox(chunk) -> None:
        """청크의 start_index + 부모 blocks_json을 이용해 bbox 유니온을 계산해 붙인다.

        결과: chunk.metadata['bboxes_json'] = '[{"page": N, "bbox": [x0,y0,x1,y1],
        "page_width": W, "page_height": H}]' (Chroma primitive 제약 때문에 문자열로).
        PDF 아닌 문서는 blocks_json이 없으므로 아무것도 하지 않음.
        """
        import json as _json
        blocks_json = chunk.metadata.get("blocks_json")
        if not blocks_json:
            return
        try:
            blocks = _json.loads(blocks_json)
        except Exception:
            return
        if not blocks:
            return
        start = chunk.metadata.get("start_index")
        if not isinstance(start, int) or start < 0:
            return
        end = start + len(chunk.page_content or "")
        # 청크 범위 [start, end)와 겹치는 블록들의 bbox 유니온 계산
        x0 = y0 = float("inf")
        x1 = y1 = float("-inf")
        matched = False
        for blk in blocks:
            b_off = blk.get("offset")
            b_len = blk.get("length")
            bb = blk.get("bbox")
            if not isinstance(b_off, int) or not isinstance(b_len, int) or not isinstance(bb, list) or len(bb) != 4:
                continue
            b_end = b_off + b_len
            # 겹침 판정
            if b_end <= start or b_off >= end:
                continue
            matched = True
            x0 = min(x0, float(bb[0]))
            y0 = min(y0, float(bb[1]))
            x1 = max(x1, float(bb[2]))
            y1 = max(y1, float(bb[3]))
        if not matched:
            return
        page_num = chunk.metadata.get("page")
        page_w = chunk.metadata.get("page_width")
        page_h = chunk.metadata.get("page_height")
        entry = {
            "page": int(page_num) if page_num is not None else None,
            "bbox": [round(x0, 2), round(y0, 2), round(x1, 2), round(y1, 2)],
        }
        if page_w is not None:
            entry["page_width"] = float(page_w)
        if page_h is not None:
            entry["page_height"] = float(page_h)
        chunk.metadata["bboxes_json"] = _json.dumps([entry], ensure_ascii=False)
        # 원본 blocks_json은 청크 메타에 계속 남아있으면 저장 용량 낭비 → 제거
        chunk.metadata.pop("blocks_json", None)

    async def process_documents(self, documents: List[Document], metadata: dict = None) -> List[Document]:
        """Async: Process documents by splitting them into chunks and adding metadata."""
        try:
            print(f"Processing {len(documents)} documents...")
            # Add additional metadata if provided
            if metadata:
                for doc in documents:
                    doc.metadata.update(metadata)
            
            # Split documents into chunks (strategy pluggable via CHUNKER_STRATEGY)
            chunks = await self.chunker.split(documents)

            # Add chunk information to metadata
            for i, chunk in enumerate(chunks):
                chunk_id = str(uuid.uuid4())
                chunk.metadata.update({
                    "chunk_id": chunk_id,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "content_length": len(chunk.page_content or ""),
                })
                # PDF 등: 0-based page가 있으면 1-based page_number 추가 (이미지 page_number와 일치)
                if "page" in chunk.metadata and chunk.metadata["page"] is not None:
                    try:
                        chunk.metadata["page_number"] = int(chunk.metadata["page"]) + 1
                    except (TypeError, ValueError):
                        pass

                # PDF bbox 역산: 부모 페이지의 blocks_json + chunk의 start_index로
                # 어떤 블록(들)이 청크에 포함됐는지 찾아 bbox 유니온을 계산.
                try:
                    self._attach_chunk_bbox(chunk)
                except Exception as bbox_err:
                    print(f"[bbox] 청크 bbox 계산 실패: {bbox_err}")

            return chunks
        except Exception as e:
            print(f"Error processing documents: {e}")
            return []

    def process_directory(self, directory_path: str, metadata: dict = None) -> List[Document]:
        """Process all documents in a directory."""
        all_documents = []
        
        for root, _, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                print(f"Processing file: {file_path}")
                
                # Read file content into memory
                with open(file_path, 'rb') as f:
                    file_content = f.read()
                
                documents = self.load_document(file_content, file_path)
                
                if documents:
                    print(f"Loaded {len(documents)} documents from {file_path}")
                    all_documents.extend(documents)
        
        if all_documents:
            print(f"Processing {len(all_documents)} documents...")
            chunks = self.process_documents(all_documents, metadata)
            print(f"Created {len(chunks)} chunks")
            return chunks
        
        return []

    async def extract_images_from_document(self, file_content: bytes, file_name: str, file_id: str) -> List[Dict[str, Any]]:
        """문서에서 이미지 추출 (PDF, DOCX, PPTX 지원) - 재사용 가능한 버전"""
        extracted_images = []
        
        print(f"Starting image extraction for file: {file_name}")
        
        # 파일 확장자 확인
        file_extension = Path(file_name).suffix.lower()
        
        if file_extension == '.pdf':
            print("Processing PDF file for image extraction...")
            extracted_images = await self._extract_images_from_pdf(file_content, file_name, file_id)
        elif file_extension == '.docx':
            print("Processing DOCX file for image extraction...")
            extracted_images = await self._extract_images_from_docx(file_content, file_name, file_id)
        elif file_extension == '.pptx':
            print("Processing PPTX file for image extraction...")
            extracted_images = await self._extract_images_from_pptx(file_content, file_name, file_id)
        else:
            print(f"File type {file_extension} not supported for image extraction")
            
        print(f"Extracted {len(extracted_images)} images from {file_name}")
        return extracted_images

    async def _extract_images_from_pdf(self, file_content: bytes, file_name: str, file_id: str) -> List[Dict[str, Any]]:
        """PDF 파일에서 이미지 추출"""
        extracted_images = []
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            try:
                doc = fitz.open(tmp_path)
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    image_list = page.get_images()
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image['image']
                            
                            # 이미지 메타데이터
                            image_metadata = {
                                'format': base_image.get('ext', 'png'),
                                'source_path': file_name,
                                'page_number': page_num + 1,
                                'image_index': img_index
                            }
                            
                            # 이미지 이름 생성
                            image_name = f"{file_id}_page{page_num+1}_img{img_index}.{image_metadata['format']}"
                            
                            extracted_images.append({
                                'image_id': f"{file_id}_page{page_num+1}_img{img_index}",
                                'image_name': image_name,
                                'image_data': image_bytes,
                                'metadata': image_metadata
                            })
                            
                        except Exception as e:
                            print(f"Error extracting image {img_index} from page {page_num + 1}: {e}")
                            continue
                
                doc.close()
                
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            print(f"Error processing PDF file {file_name}: {e}")
            
        return extracted_images

    async def _extract_images_from_docx(self, file_content: bytes, file_name: str, file_id: str) -> List[Dict[str, Any]]:
        """DOCX 파일에서 이미지 추출"""
        extracted_images = []
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name
            
            try:
                with zipfile.ZipFile(tmp_path, 'r') as docx_zip:
                    image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                    
                    for img_index, image_path in enumerate(image_files):
                        try:
                            # 이미지 데이터 추출
                            image_data = docx_zip.read(image_path)
                            
                            # 파일 확장자 추출
                            file_extension = Path(image_path).suffix.lower()
                            if not file_extension:
                                # MIME 타입으로 판단
                                if image_data.startswith(b'\xff\xd8\xff'):
                                    file_extension = '.jpg'
                                elif image_data.startswith(b'\x89PNG\r\n\x1a\n'):
                                    file_extension = '.png'
                                elif image_data.startswith(b'GIF8'):
                                    file_extension = '.gif'
                                else:
                                    file_extension = '.png'
                            
                            # 이미지 이름 생성
                            image_name = f"{file_id}_img{img_index+1}{file_extension}"
                            
                            # 이미지 메타데이터 생성
                            image_metadata = {
                                'format': file_extension[1:],
                                'source_path': image_path,
                                'image_index': img_index + 1
                            }
                            
                            extracted_images.append({
                                'image_id': f"{file_id}_img{img_index+1}",
                                'image_name': image_name,
                                'image_data': image_data,
                                'metadata': image_metadata
                            })
                            
                        except Exception as e:
                            print(f"Error extracting image {image_path}: {e}")
                            continue
                            
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            print(f"Error processing DOCX file {file_name}: {e}")
            
        return extracted_images

    async def _extract_images_from_pptx(self, file_content: bytes, file_name: str, file_id: str) -> List[Dict[str, Any]]:
        """PPTX 파일에서 이미지 추출"""
        extracted_images = []
        
        try:
            print(f"Starting PPTX image extraction for {file_name}")
            
            # 임시 파일로 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pptx') as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name
            
            try:
                # PPTX는 ZIP 파일이므로 압축 해제
                with zipfile.ZipFile(tmp_path, 'r') as pptx_zip:
                    # 이미지 파일들 찾기
                    image_files = [f for f in pptx_zip.namelist() if f.startswith('ppt/media/')]
                    print(f"Found {len(image_files)} image files in PPTX")
                    
                    for img_index, image_path in enumerate(image_files):
                        try:
                            print(f"Processing image {img_index + 1}: {image_path}")
                            
                            # 이미지 데이터 추출
                            image_data = pptx_zip.read(image_path)
                            
                            # 파일 확장자 추출
                            file_extension = Path(image_path).suffix.lower()
                            if not file_extension:
                                # 확장자가 없는 경우 MIME 타입으로 판단
                                if image_data.startswith(b'\xff\xd8\xff'):
                                    file_extension = '.jpg'
                                elif image_data.startswith(b'\x89PNG\r\n\x1a\n'):
                                    file_extension = '.png'
                                elif image_data.startswith(b'GIF8'):
                                    file_extension = '.gif'
                                else:
                                    file_extension = '.png'  # 기본값
                            
                            print(f"Detected image format: {file_extension}")
                            
                            # 이미지 이름 생성
                            image_name = f"{file_id}_slide_img{img_index+1}{file_extension}"
                            
                            # 이미지 메타데이터 생성
                            image_metadata = {
                                'format': file_extension[1:],  # 확장자에서 점 제거
                                'source_path': image_path,
                                'slide_index': img_index + 1
                            }
                            
                            extracted_images.append({
                                'image_id': f"{file_id}_slide_img{img_index+1}",
                                'image_name': image_name,
                                'image_data': image_data,
                                'metadata': image_metadata
                            })
                            
                            print(f"Successfully extracted image {img_index + 1}")
                            
                        except Exception as e:
                            print(f"Error extracting image {image_path}: {e}")
                            continue
                            
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            print(f"Error processing PPTX file {file_name}: {e}")
            
        return extracted_images

    async def iter_extract_images_from_document(
        self, file_content: bytes, file_name: str, file_id: str
    ) -> AsyncIterator[Dict[str, Any]]:
        """Yield one image at a time so callers can process in batches without holding all bytes."""
        file_extension = Path(file_name).suffix.lower()
        if file_extension == '.pdf':
            async for img in self._iter_pdf_images(file_content, file_name, file_id):
                yield img
        elif file_extension == '.docx':
            async for img in self._iter_zip_images(file_content, file_id, suffix='.docx',
                                                   prefix='word/media/', id_tpl='{file_id}_img{i}'):
                yield img
        elif file_extension == '.pptx':
            async for img in self._iter_zip_images(file_content, file_id, suffix='.pptx',
                                                   prefix='ppt/media/', id_tpl='{file_id}_slide_img{i}'):
                yield img

    async def _iter_pdf_images(
        self, file_content: bytes, file_name: str, file_id: str
    ) -> AsyncIterator[Dict[str, Any]]:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name
        try:
            doc = fitz.open(tmp_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                for img_index, img in enumerate(page.get_images()):
                    try:
                        base_image = doc.extract_image(img[0])
                        image_bytes = base_image['image']
                        ext = base_image.get('ext', 'png')
                        yield {
                            'image_id': f"{file_id}_page{page_num+1}_img{img_index}",
                            'image_name': f"{file_id}_page{page_num+1}_img{img_index}.{ext}",
                            'image_data': image_bytes,
                            'metadata': {
                                'format': ext,
                                'source_path': file_name,
                                'page_number': page_num + 1,
                                'image_index': img_index,
                            },
                        }
                    except Exception as e:
                        print(f"Error extracting image {img_index} from page {page_num + 1}: {e}")
                        continue
            doc.close()
        finally:
            os.unlink(tmp_path)

    async def _iter_zip_images(
        self, file_content: bytes, file_id: str, suffix: str, prefix: str, id_tpl: str
    ) -> AsyncIterator[Dict[str, Any]]:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name
        try:
            with zipfile.ZipFile(tmp_path, 'r') as zf:
                image_files = [f for f in zf.namelist() if f.startswith(prefix)]
                for i, image_path in enumerate(image_files, start=1):
                    try:
                        data = zf.read(image_path)
                        ext = Path(image_path).suffix.lower()
                        if not ext:
                            if data.startswith(b'\xff\xd8\xff'):
                                ext = '.jpg'
                            elif data.startswith(b'\x89PNG\r\n\x1a\n'):
                                ext = '.png'
                            elif data.startswith(b'GIF8'):
                                ext = '.gif'
                            else:
                                ext = '.png'
                        image_id = id_tpl.format(file_id=file_id, i=i)
                        yield {
                            'image_id': image_id,
                            'image_name': f"{image_id}{ext}",
                            'image_data': data,
                            'metadata': {
                                'format': ext[1:],
                                'source_path': image_path,
                                'image_index': i,
                            },
                        }
                    except Exception as e:
                        print(f"Error extracting image {image_path}: {e}")
                        continue
        finally:
            os.unlink(tmp_path)

    async def extract_and_upload_images_batched(
        self,
        file_content: bytes,
        file_name: str,
        file_id: str,
        tenant_id: str,
        batch_size: int = 15,
    ) -> List[Dict[str, Any]]:
        """Extract and upload images in batches; raw bytes released between batches."""
        from app.core import config
        if not config.image_analysis_enabled():
            print("Skipping image extraction/upload (vision disabled)")
            return []

        from app.storage.image_storage import get_image_storage_utils
        storage = get_image_storage_utils()
        folder = f"extracted_images/{tenant_id}/{file_id}"

        uploaded: List[Dict[str, Any]] = []
        batch: List[Dict[str, Any]] = []
        total = 0

        async def flush() -> None:
            nonlocal total
            if not batch:
                return
            results = await asyncio.gather(*[
                storage.upload_image_to_storage(img['image_data'], img['image_name'], folder)
                for img in batch
            ], return_exceptions=True)
            for img, result in zip(batch, results):
                if isinstance(result, Exception) or not result:
                    print(f"Error uploading {img['image_name']}: {result}")
                    continue
                uploaded.append({
                    'image_id': img['image_id'],
                    'image_name': img['image_name'],
                    'image_url': result.get('public_url'),
                    'metadata': img['metadata'],
                })
            total += len(batch)
            print(f"Uploaded {total} images (batch of {len(batch)})")
            batch.clear()

        async for img in self.iter_extract_images_from_document(file_content, file_name, file_id):
            batch.append(img)
            if len(batch) >= batch_size:
                await flush()
        await flush()
        return uploaded


_document_processor_instance: Optional["DocumentProcessor"] = None


def get_document_processor() -> "DocumentProcessor":
    global _document_processor_instance
    if _document_processor_instance is None:
        _document_processor_instance = DocumentProcessor()
    return _document_processor_instance


# Example usage
if __name__ == "__main__":
    processor = DocumentProcessor()
    documents = processor.process_directory("./documents")
    print(f"Processed {len(documents)} document chunks")
