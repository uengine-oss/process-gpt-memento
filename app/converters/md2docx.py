#!/usr/bin/env python3
"""
Markdown to DOCX converter with Mermaid diagrams as images
Usage: python3 md2docx_converter.py <input.md> [--output output.docx] [--image-format FORMAT]
"""

import argparse
import subprocess
import tempfile
import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from io import BytesIO

def generate_mermaid_image(mermaid_code, output_path, image_format='png'):
    """
    Generate image from Mermaid code using mermaid-cli
    Returns True if successful, False otherwise
    """
    try:
        # Create a temporary file for the mermaid code
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as temp_mmd:
            temp_mmd.write(mermaid_code)
            temp_mmd_path = temp_mmd.name
        
        # Run mermaid-cli to generate image
        result = subprocess.run([
            'mmdc', '-i', temp_mmd_path, '-o', output_path,
            '--backgroundColor', 'white',
            '--width', '1200',
            '--height', '800'
        ], capture_output=True, text=True)
        
        # Clean up temporary file
        os.unlink(temp_mmd_path)
        
        if result.returncode == 0 and os.path.exists(output_path):
            return True
        else:
            print(f"Error generating Mermaid image: {result.stderr}")
            return False
            
    except Exception as e:
        print(f"Exception in Mermaid image generation: {e}")
        return False

def process_markdown_with_mermaid(md_content, temp_dir, image_format='png'):
    """
    Process Markdown content and replace Mermaid blocks with image references
    Returns processed markdown and list of generated image files
    """
    image_files = []
    image_counter = 1
    
    def replace_mermaid_block(match):
        nonlocal image_counter
        mermaid_code = match.group(1).strip()
        
        # Generate image filename
        image_filename = f"mermaid_diagram_{image_counter}.{image_format}"
        image_path = os.path.join(temp_dir, image_filename)
        
        # Generate image
        if generate_mermaid_image(mermaid_code, image_path, image_format):
            image_files.append(image_path)
            image_counter += 1
            # Return markdown image syntax
            return f"![Mermaid Diagram {image_counter-1}]({image_path})"
        else:
            # Return error message if image generation failed
            return f"**[Mermaid Diagram Error: Could not generate image]**"
    
    # Replace all mermaid code blocks
    processed_md = re.sub(r'```mermaid\n(.*?)\n```', replace_mermaid_block, md_content, flags=re.DOTALL)
    
    return processed_md, image_files

def convert_md_to_docx_with_pandoc(md_content, output_path, temp_dir):
    """
    Convert processed Markdown to DOCX using pandoc
    """
    try:
        # Create temporary markdown file
        temp_md_path = os.path.join(temp_dir, 'temp_processed.md')
        with open(temp_md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        # Use pandoc to convert to DOCX
        result = subprocess.run([
            'pandoc', 
            temp_md_path,
            '-o', output_path,
            '--from', 'markdown',
            '--to', 'docx',
            '--standalone',
            '--toc',  # Table of contents
            '--reference-doc', '/dev/null'  # Use default styling
        ], capture_output=True, text=True)
        
        if result.returncode == 0:
            return True
        else:
            print(f"Pandoc error: {result.stderr}")
            return False
            
    except Exception as e:
        print(f"Exception in pandoc conversion: {e}")
        return False

def create_docx_manually(md_content, output_path):
    """
    Create DOCX manually using python-docx (fallback method)
    """
    try:
        # Parse markdown to HTML first
        md = markdown.Markdown(extensions=[
            'extra',
            'codehilite',
            'toc',
            'fenced_code',
            'tables'
        ])
        
        html_content = md.convert(md_content)
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('대규모 딥러닝 시스템 아키텍처 설계 문서', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Process the markdown content line by line
        lines = md_content.split('\n')
        current_paragraph = None
        in_code_block = False
        
        for line in lines:
            line = line.strip()
            
            if not line and not in_code_block:
                current_paragraph = None
                continue
            
            # Handle headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                doc.add_heading(heading_text, level)
                current_paragraph = None
            
            # Handle code blocks
            elif line.startswith('```'):
                in_code_block = not in_code_block
                current_paragraph = None
            
            # Handle images (including our generated Mermaid images)
            elif line.startswith('!['):
                image_match = re.match(r'!\[.*?\]\((.*?)\)', line)
                if image_match:
                    image_path = image_match.group(1)
                    if os.path.exists(image_path):
                        try:
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                            run.add_picture(image_path, width=Inches(6))
                        except Exception as e:
                            doc.add_paragraph(f"[Image could not be inserted: {image_path}]")
                    current_paragraph = None
            
            # Handle regular text
            elif not in_code_block:
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                
                # Handle bold text
                if '**' in line:
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            current_paragraph.add_run(part)
                        else:
                            current_paragraph.add_run(part).bold = True
                else:
                    current_paragraph.add_run(line + ' ')
        
        # Save document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Exception in manual DOCX creation: {e}")
        return False

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown files with Mermaid diagrams to DOCX",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Image formats:
  png    - PNG format (default, good quality, larger file size)
  svg    - SVG format (vector graphics, scalable)
  pdf    - PDF format (vector graphics, high quality)

Examples:
  python3 md2docx_converter.py sample.md
  python3 md2docx_converter.py sample.md --output report.docx --image-format png
  python3 md2docx_converter.py sample.md --image-format svg
        """
    )
    
    parser.add_argument('input', help='Input Markdown file')
    parser.add_argument('--output', '-o', help='Output DOCX file (default: input filename with .docx extension)')
    parser.add_argument('--image-format', '-f', 
                       choices=['png', 'svg', 'pdf'],
                       default='png',
                       help='Image format for Mermaid diagrams (default: png)')
    
    args = parser.parse_args()
    
    # Validate input file
    if not os.path.exists(args.input):
        print(f"Error: Input file '{args.input}' not found!")
        sys.exit(1)
    
    # Generate output filename if not provided
    if args.output is None:
        base_name = os.path.splitext(args.input)[0]
        args.output = f"{base_name}.docx"
    
    # Create temporary directory for images
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            print(f"Converting {args.input} to {args.output}...")
            print(f"Image format: {args.image_format}")
            
            # Read input file
            with open(args.input, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Process Mermaid diagrams
            print("Processing Mermaid diagrams...")
            processed_md, image_files = process_markdown_with_mermaid(
                md_content, temp_dir, args.image_format
            )
            
            print(f"Generated {len(image_files)} diagram images")
            
            # Try pandoc conversion first
            print("Converting to DOCX using pandoc...")
            if convert_md_to_docx_with_pandoc(processed_md, args.output, temp_dir):
                print("✅ Conversion completed successfully using pandoc!")
            else:
                print("⚠️  Pandoc conversion failed, trying manual conversion...")
                if create_docx_manually(processed_md, args.output):
                    print("✅ Conversion completed successfully using manual method!")
                else:
                    print("❌ All conversion methods failed!")
                    sys.exit(1)
            
            # Show file size
            file_size = os.path.getsize(args.output)
            print(f"📁 Output file: {args.output}")
            print(f"📊 File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
            print(f"🖼️  Mermaid diagrams: {len(image_files)} images included")
            
        except Exception as e:
            print(f"❌ Error during conversion: {e}")
            sys.exit(1)

if __name__ == "__main__":
    main()
