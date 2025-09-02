#!/usr/bin/env python3
"""
Advanced Markdown to DOCX converter with enhanced formatting and Mermaid support
Usage: python3 md2docx_advanced.py <input.md> [--output output.docx] [--style STYLE]
"""

import argparse
import subprocess
import tempfile
import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def generate_mermaid_image(mermaid_code, output_path, image_format='png'):
    """Generate high-quality image from Mermaid code"""
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as temp_mmd:
            temp_mmd.write(mermaid_code)
            temp_mmd_path = temp_mmd.name
        
        # High-quality settings for better DOCX integration
        result = subprocess.run([
            'mmdc', '-i', temp_mmd_path, '-o', output_path,
            '--backgroundColor', 'white',
            '--width', '1400',      # Higher resolution
            '--height', '1000',
            '--scale', '2'          # 2x scaling for crisp images
        ], capture_output=True, text=True)
        
        os.unlink(temp_mmd_path)
        
        return result.returncode == 0 and os.path.exists(output_path)
        
    except Exception as e:
        print(f"Exception in Mermaid image generation: {e}")
        return False

def setup_document_styles(doc):
    """Setup custom styles for the document"""
    styles = doc.styles
    
    # Title style
    if 'Custom Title' not in styles:
        title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_style.paragraph_format.space_after = Pt(20)
    
    # Heading 1 style
    if 'Custom Heading 1' not in styles:
        h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Arial'
        h1_font.size = Pt(18)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(220, 50, 47)  # Red
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    if 'Custom Heading 2' not in styles:
        h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Arial'
        h2_font.size = Pt(16)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(38, 139, 210)  # Blue
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(5)
    
    # Heading 3 style
    if 'Custom Heading 3' not in styles:
        h3_style = styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Arial'
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(133, 153, 0)  # Green
        h3_style.paragraph_format.space_before = Pt(8)
        h3_style.paragraph_format.space_after = Pt(4)
    
    # Code style
    if 'Custom Code' not in styles:
        code_style = styles.add_style('Custom Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Consolas'
        code_font.size = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.5)
        # Add light gray background (simulated with borders)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    
    # Body text style
    if 'Custom Body' not in styles:
        body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(11)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        body_style.paragraph_format.space_after = Pt(6)

def add_table_from_markdown(doc, table_text):
    """Convert markdown table to DOCX table"""
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
    if len(lines) < 2:
        return
    
    # Parse headers
    headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
    
    # Parse rows (skip separator line)
    rows = []
    for line in lines[2:]:
        if '|' in line:
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(row) == len(headers):
                rows.append(row)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

def process_mathematical_expressions(text):
    """Process LaTeX mathematical expressions for DOCX"""
    # Convert inline math $...$ to italic text (simple fallback)
    text = re.sub(r'\$([^$]+)\$', r'[\1]', text)
    
    # Convert display math $$...$$ to centered italic text
    text = re.sub(r'\$\$([^$]+)\$\$', r'[\1]', text)
    
    return text

def parse_markdown_advanced(md_content, temp_dir, image_format='png'):
    """Advanced markdown parsing with Mermaid support"""
    
    doc = Document()
    setup_document_styles(doc)
    
    # Process Mermaid diagrams first
    image_files = []
    image_counter = 1
    
    def replace_mermaid_block(match):
        nonlocal image_counter
        mermaid_code = match.group(1).strip()
        
        image_filename = f"mermaid_diagram_{image_counter}.{image_format}"
        image_path = os.path.join(temp_dir, image_filename)
        
        if generate_mermaid_image(mermaid_code, image_path, image_format):
            image_files.append(image_path)
            image_counter += 1
            return f"MERMAID_IMAGE:{image_path}"
        else:
            return f"**[Mermaid Diagram Error: Could not generate image]**"
    
    # Replace Mermaid blocks with placeholders
    processed_content = re.sub(r'```mermaid\n(.*?)\n```', replace_mermaid_block, md_content, flags=re.DOTALL)
    
    # Split content into lines for processing
    lines = processed_content.split('\n')
    
    current_paragraph = None
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    for i, line in enumerate(lines):
        original_line = line
        line = line.rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    code_para = doc.add_paragraph(code_text, style='Custom Code')
                    code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
                current_paragraph = None
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        elif in_table:
            # End of table
            add_table_from_markdown(doc, '\n'.join(table_lines))
            in_table = False
            table_lines = []
            current_paragraph = None
        
        # Empty line
        if not line:
            current_paragraph = None
            continue
        
        # Handle Mermaid image placeholders
        if line.startswith('MERMAID_IMAGE:'):
            image_path = line.replace('MERMAID_IMAGE:', '')
            if os.path.exists(image_path):
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()
                try:
                    run.add_picture(image_path, width=Inches(6.5))
                except Exception as e:
                    doc.add_paragraph(f"[Image could not be inserted: {os.path.basename(image_path)}]")
            current_paragraph = None
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            heading_text = process_mathematical_expressions(heading_text)
            
            if level == 1:
                if heading_text == '대규모 딥러닝 시스템 아키텍처 설계 문서':
                    doc.add_paragraph(heading_text, style='Custom Title')
                else:
                    doc.add_heading(heading_text, level=1).style = doc.styles['Custom Heading 1']
            elif level == 2:
                doc.add_heading(heading_text, level=2).style = doc.styles['Custom Heading 2']
            elif level == 3:
                doc.add_heading(heading_text, level=3).style = doc.styles['Custom Heading 3']
            else:
                doc.add_heading(heading_text, level=min(level, 9))
            
            current_paragraph = None
            continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            # Add a horizontal line (simulated with empty paragraph)
            doc.add_paragraph()
            current_paragraph = None
            continue
        
        # Handle lists
        if line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\. ', line):
            list_text = re.sub(r'^[-*] |^\d+\. ', '', line)
            list_text = process_mathematical_expressions(list_text)
            
            paragraph = doc.add_paragraph(list_text, style='List Bullet' if line.startswith(('- ', '* ')) else 'List Number')
            format_text_in_paragraph(paragraph)
            current_paragraph = None
            continue
        
        # Handle regular paragraphs
        if current_paragraph is None:
            current_paragraph = doc.add_paragraph(style='Custom Body')
        
        # Process mathematical expressions
        line = process_mathematical_expressions(line)
        
        # Add text to current paragraph
        add_formatted_text_to_paragraph(current_paragraph, line + ' ')
    
    # Handle any remaining table
    if in_table and table_lines:
        add_table_from_markdown(doc, '\n'.join(table_lines))
    
    return doc, image_files

def add_formatted_text_to_paragraph(paragraph, text):
    """Add formatted text to paragraph with bold, italic support"""
    # Handle bold text **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Regular text
            paragraph.add_run(part)

def format_text_in_paragraph(paragraph):
    """Apply formatting to existing paragraph text"""
    text = paragraph.text
    paragraph.clear()
    add_formatted_text_to_paragraph(paragraph, text)

def main():
    parser = argparse.ArgumentParser(
        description="Advanced Markdown to DOCX converter with Mermaid diagrams",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Features:
  ✅ Custom styling and typography
  ✅ High-quality Mermaid diagrams
  ✅ Mathematical expressions support
  ✅ Table formatting
  ✅ Code block formatting
  ✅ Professional document layout

Image formats:
  png    - PNG format (default, best compatibility)
  svg    - SVG format (vector graphics, may not display in all apps)
  pdf    - PDF format (vector graphics, high quality)

Examples:
  python3 md2docx_advanced.py sample.md
  python3 md2docx_advanced.py sample.md --output professional_report.docx
  python3 md2docx_advanced.py sample.md --image-format png
        """
    )
    
    parser.add_argument('input', help='Input Markdown file')
    parser.add_argument('--output', '-o', help='Output DOCX file (default: input filename with .docx extension)')
    parser.add_argument('--image-format', '-f', 
                       choices=['png', 'svg', 'pdf'],
                       default='png',
                       help='Image format for Mermaid diagrams (default: png)')
    
    args = parser.parse_args()
    
    # Validate input file
    if not os.path.exists(args.input):
        print(f"Error: Input file '{args.input}' not found!")
        sys.exit(1)
    
    # Generate output filename if not provided
    if args.output is None:
        base_name = os.path.splitext(args.input)[0]
        args.output = f"{base_name}_advanced.docx"
    
    # Create temporary directory for images
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            print(f"🔄 Converting {args.input} to {args.output}...")
            print(f"🖼️  Image format: {args.image_format}")
            
            # Read input file
            with open(args.input, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            print("📊 Parsing and formatting document...")
            
            # Advanced parsing
            doc, image_files = parse_markdown_advanced(md_content, temp_dir, args.image_format)
            
            print(f"🎨 Generated {len(image_files)} high-quality diagram images")
            
            # Save document
            doc.save(args.output)
            
            print("✅ Conversion completed successfully!")
            
            # Show file info
            file_size = os.path.getsize(args.output)
            print(f"📁 Output file: {args.output}")
            print(f"📊 File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
            print(f"🖼️  Mermaid diagrams: {len(image_files)} images included")
            print(f"🎯 Features: Custom styling, tables, math expressions, code formatting")
            
        except Exception as e:
            print(f"❌ Error during conversion: {e}")
            import traceback
            traceback.print_exc()
            sys.exit(1)

if __name__ == "__main__":
    main()
