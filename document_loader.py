"""
Document loader and processor
"""
import os
import re
import docx
import uuid
import tempfile
import asyncio
import io
import zipfile
from typing import List, Optional, Dict, Any, Tuple
from pathlib import Path
from pydantic import BaseModel
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from openai import OpenAI
from langchain_community.document_loaders import (
    UnstructuredWordDocumentLoader,
    UnstructuredPowerPointLoader,
    UnstructuredExcelLoader,
    UnstructuredFileLoader,
    PyPDFLoader,
    TextLoader
)
import fitz  # PyMuPDF for PDF image extraction

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# Allow loading vendored extract_hwp when the installed extract-hwp package has no module (PyPI 0.1.0 packaging bug)
_vendor_dir = Path(__file__).resolve().parent / "vendor"
if _vendor_dir.is_dir() and str(_vendor_dir) not in __import__("sys").path:
    __import__("sys").path.insert(0, str(_vendor_dir))


def _extract_text_from_hwp_or_hwpx(file_path: str, file_extension: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text from HWP or HWPX file. Uses extract_hwp if available; otherwise .hwpx via zip+xml.
    PyPI extract-hwp 0.1.0 ships no module (packaging bug). For .hwp either install from Git or add
    vendor: clone https://github.com/thlee/extract-hwp and copy src/extract_hwp into vendor/.
    Returns (text, error_message); error_message is None on success.
    """
    try:
        from extract_hwp import extract_text_from_hwp
        return extract_text_from_hwp(file_path)
    except ModuleNotFoundError:
        pass  # extract-hwp not installed or broken (PyPI 0.1.0 has no module); use fallback

    if file_extension == ".hwp":
        return (
            None,
            "HWP 파일 처리를 위해 extract-hwp를 GitHub에서 설치해 주세요: uv pip install \"extract-hwp @ git+https://github.com/thlee/extract-hwp.git\" (또는 vendor 폴더에 소스 추가)",
        )

    if file_extension == ".hwpx":
        try:
            import xml.etree.ElementTree as ET
            with zipfile.ZipFile(file_path, "r") as z:
                names = z.namelist()
                section_files = sorted(n for n in names if "Contents/section" in n and n.endswith(".xml"))
                if not section_files:
                    content_name = next((n for n in names if "contents" in n.lower() and n.endswith(".xml")), None)
                    if not content_name:
                        return (None, "HWPX: section or contents XML not found")
                    with z.open(content_name) as f:
                        raw = f.read().decode("utf-8", errors="replace")
                    text = re.sub(r"<[^>]+>", " ", raw)
                    text = re.sub(r"\s+", " ", text).strip()
                    return (text, None)
                sections = []
                for section_file in section_files:
                    with z.open(section_file) as f:
                        root = ET.fromstring(f.read())
                    lines = []
                    for elem in root.iter():
                        if elem.tag.endswith("}p"):
                            line = "".join(
                                (n.text or "") for n in elem.iter() if n.tag.endswith("}t")
                            )
                            if line.strip():
                                lines.append(line)
                    if lines:
                        sections.append("\n".join(lines))
                return ("\n\n".join(sections), None)
        except Exception as e:
            return (None, str(e))

    return (None, f"Unsupported extension: {file_extension}")


class DocumentProcessor:
    def __init__(self, chunk_size: int = 2000, chunk_overlap: int = 400):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", " ", ""],
            length_function=len,
            is_separator_regex=False
        )
        self._openai_client: Optional[OpenAI] = None

    def _get_openai_client(self) -> OpenAI:
        if self._openai_client is None:
            self._openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        return self._openai_client

    async def _generate_section_title_single(
        self, content: str, semaphore: asyncio.Semaphore
    ) -> str:
        """청크 1개의 section_title을 Structured Output으로 생성한다."""

        class _TitleResponse(BaseModel):
            title: str

        snippet = content[:300].replace("\n", " ")
        async with semaphore:
            try:
                client = self._get_openai_client()
                response = await asyncio.to_thread(
                    client.beta.chat.completions.parse,
                    model="gpt-4o-mini",
                    messages=[{
                        "role": "user",
                        "content": (
                            "다음 문서 내용에 어울리는 소제목(10자 이내)을 생성하세요.\n\n"
                            f"{snippet}"
                        ),
                    }],
                    response_format=_TitleResponse,
                    temperature=0,
                    max_tokens=50,
                )
                parsed = response.choices[0].message.parsed
                return parsed.title.strip() if parsed else ""
            except Exception as e:
                print(f"section_title 생성 실패: {e}")
                return ""

    async def _generate_section_titles(self, chunks: List[Document]) -> List[str]:
        """청크별 section_title을 병렬로 생성한다 (Structured Output 사용).

        동시 요청 수를 semaphore로 제한해 rate limit를 방지한다.
        """
        semaphore = asyncio.Semaphore(10)
        tasks = [
            self._generate_section_title_single(chunk.page_content or "", semaphore)
            for chunk in chunks
        ]
        return list(await asyncio.gather(*tasks))

    def _load_docx_with_python_docx(self, tmp_path: str, file_name: str) -> List[Document]:
        """DOCX 텍스트 추출 paragraph + table 셀 텍스트."""
        if docx is None:
            raise RuntimeError("python-docx is not installed")
        doc = docx.Document(tmp_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = "\n\n".join(parts) if parts else ""
        return [Document(page_content=text)]

    @staticmethod
    def _table_to_markdown(table: List[List[Any]]) -> str:
        """표를 마크다운 형식으로 변환"""
        if not table or len(table) == 0:
            return ""
        markdown = []
        header = table[0]
        markdown.append("| " + " | ".join(str(cell) if cell is not None and cell != "" else "" for cell in header) + " |")
        markdown.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in table[1:]:
            markdown.append("| " + " | ".join(str(cell) if cell is not None and cell != "" else "" for cell in row) + " |")
        return "\n".join(markdown)

    @staticmethod
    def _get_text_blocks_outside_tables(page: Any, table_bboxes: List[tuple]) -> List[Dict[str, Any]]:
        """표 영역을 제외하고 y좌표별 텍스트 블록 추출 (pdfplumber Page 사용)"""
        def word_in_bbox(word: Dict, bbox: tuple) -> bool:
            x0, top, x1, bottom = bbox
            h_mid = (word["x0"] + word["x1"]) / 2
            v_mid = (word["top"] + word["bottom"]) / 2
            return x0 <= h_mid <= x1 and top <= v_mid <= bottom

        try:
            all_words = page.extract_words()
        except Exception:
            return []
        words = [w for w in all_words if not any(word_in_bbox(w, bbox) for bbox in table_bboxes)]
        if not words:
            return []

        lines: Dict[float, List[tuple]] = {}
        for word in words:
            y = round(word["top"], 1)
            if y not in lines:
                lines[y] = []
            lines[y].append((word["x0"], word.get("text", "")))

        text_blocks = []
        current_block: List[tuple] = []
        prev_y = None
        for y in sorted(lines.keys()):
            line_words = sorted(lines[y], key=lambda x: x[0])
            line_text = " ".join(w[1] for w in line_words)
            if prev_y is not None and (y - prev_y) > 20:
                if current_block:
                    text_blocks.append({
                        "type": "text",
                        "y": current_block[0][0],
                        "content": "\n".join(line[1] for line in current_block),
                    })
                current_block = []
            current_block.append((y, line_text))
            prev_y = y
        if current_block:
            text_blocks.append({
                "type": "text",
                "y": current_block[0][0],
                "content": "\n".join(line[1] for line in current_block),
            })
        return text_blocks

    def _load_pdf_with_pdfplumber(self, tmp_path: str, file_name: str) -> List[Document]:
        """pdfplumber로 PDF 파싱: 표, 텍스트, 이미지 플레이스홀더를 y좌표 순서대로 결합"""
        documents = []
        with pdfplumber.open(tmp_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                elements: List[Dict[str, Any]] = []
                tables = page.find_tables()
                for table in tables:
                    try:
                        extracted = table.extract()
                        if extracted:
                            elements.append({
                                "type": "table",
                                "y": table.bbox[1],
                                "content": extracted,
                            })
                    except Exception:
                        continue
                table_bboxes = [t.bbox for t in tables]
                text_blocks = self._get_text_blocks_outside_tables(page, table_bboxes)
                elements.extend(text_blocks)

                # 이미지 플레이스홀더를 y좌표 기준으로 텍스트/표 사이에 삽입
                try:
                    for img_index, img in enumerate(page.images):
                        top = img.get("top")
                        if top is not None:
                            elements.append({
                                "type": "image_placeholder",
                                "y": top,
                                "image_index": img_index,
                                "page_num_1based": page_num + 1,
                            })
                except Exception:
                    pass

                elements.sort(key=lambda x: x["y"])

                parts = []
                for elem in elements:
                    if elem["type"] == "text":
                        if elem.get("content"):
                            parts.append(elem["content"])
                    elif elem["type"] == "table":
                        md = self._table_to_markdown(elem.get("content"))
                        if md:
                            parts.append(md)
                    elif elem["type"] == "image_placeholder":
                        parts.append(
                            f"__IMAGE_PLACEHOLDER_p{elem['page_num_1based']}_i{elem['image_index']}__"
                        )

                page_content = "\n\n".join(parts) if parts else ""
                doc = Document(
                    page_content=page_content,
                    metadata={
                        "source": file_name,
                        "page": page_num,
                    },
                )
                documents.append(doc)
        return documents

    async def load_document(self, file_content: bytes, file_name: str) -> Optional[List[Document]]:
        """Async: Load a document from memory (BytesIO object)."""
        try:
            file_extension = os.path.splitext(file_name)[1].lower()
            
            if file_extension == '.txt':
                content = await asyncio.to_thread(file_content.read)
                content = content.decode('utf-8-sig')
                documents = [Document(page_content=content)]
            elif file_extension == '.docx':
                # Save BytesIO to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    await asyncio.to_thread(tmp.write, file_content.read())
                    tmp_path = tmp.name
                try:
                    documents = await asyncio.to_thread(
                        self._load_docx_with_python_docx,
                        tmp_path,
                        file_name
                    )
                finally:
                    await asyncio.to_thread(os.unlink, tmp_path)
            elif file_extension == '.pptx':
                # Save BytesIO to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pptx') as tmp:
                    await asyncio.to_thread(tmp.write, file_content.read())
                    tmp_path = tmp.name
                try:
                    loader = UnstructuredPowerPointLoader(tmp_path, mode="single")
                    documents = await asyncio.to_thread(loader.load)
                finally:
                    await asyncio.to_thread(os.unlink, tmp_path)
            elif file_extension == '.xlsx':
                # Save BytesIO to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
                    await asyncio.to_thread(tmp.write, file_content.read())
                    tmp_path = tmp.name
                try:
                    loader = UnstructuredExcelLoader(tmp_path, mode="single")
                    documents = await asyncio.to_thread(loader.load)
                finally:
                    await asyncio.to_thread(os.unlink, tmp_path)
            elif file_extension == '.pdf':
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                    await asyncio.to_thread(tmp.write, file_content.read())
                    tmp_path = tmp.name
                try:
                    documents = await asyncio.to_thread(
                        self._load_pdf_with_pdfplumber, tmp_path, file_name
                    )
                finally:
                    await asyncio.to_thread(os.unlink, tmp_path)
            elif file_extension in ('.hwp', '.hwpx'):
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp:
                    await asyncio.to_thread(tmp.write, file_content.read())
                    tmp_path = tmp.name
                try:
                    text, error = await asyncio.to_thread(
                        _extract_text_from_hwp_or_hwpx, tmp_path, file_extension
                    )
                    if error is not None:
                        print(f"HWP/HWPX extraction error for {file_name}: {error}")
                        return None
                    documents = [Document(page_content=text or "")]
                finally:
                    await asyncio.to_thread(os.unlink, tmp_path)
            else:
                print(f"Unsupported file type: {file_extension}")
                return None

            # Add basic metadata including UUID
            doc_id = str(uuid.uuid4())  # Generate a single UUID for the document
            for doc in documents:
                doc.metadata.update({
                    "id": doc_id,
                    "source": file_name,
                    "file_name": file_name,
                    "file_type": file_extension[1:],
                    "language": "ko",
                    "content_length": len(doc.page_content)
                })
            
            return documents
        except Exception as e:
            print(f"Error loading document from memory {file_name}: {e}")
            return None

    async def process_documents(self, documents: List[Document], metadata: dict = None) -> List[Document]:
        """Async: Process documents by splitting them into chunks and adding metadata."""
        try:
            print(f"Processing {len(documents)} documents...")
            # Add additional metadata if provided
            if metadata:
                for doc in documents:
                    doc.metadata.update(metadata)
            
            # Split documents into chunks
            chunks = await asyncio.to_thread(self.text_splitter.split_documents, documents)

            # Add chunk information to metadata
            for i, chunk in enumerate(chunks):
                chunk_id = str(uuid.uuid4())
                chunk.metadata.update({
                    "chunk_id": chunk_id,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "content_length": len(chunk.page_content or ""),
                })
                # PDF 등: 0-based page가 있으면 1-based page_number 추가 (이미지 page_number와 일치)
                if "page" in chunk.metadata and chunk.metadata["page"] is not None:
                    try:
                        chunk.metadata["page_number"] = int(chunk.metadata["page"]) + 1
                    except (TypeError, ValueError):
                        pass

            # LLM으로 각 청크에 section_title 생성
            print(f"Generating section_title for {len(chunks)} chunks...")
            section_titles = await self._generate_section_titles(chunks)
            for chunk, title in zip(chunks, section_titles):
                if title:
                    chunk.metadata["section_title"] = title

            return chunks
        except Exception as e:
            print(f"Error processing documents: {e}")
            return []

    def process_directory(self, directory_path: str, metadata: dict = None) -> List[Document]:
        """Process all documents in a directory."""
        all_documents = []
        
        for root, _, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                print(f"Processing file: {file_path}")
                
                # Read file content into memory
                with open(file_path, 'rb') as f:
                    file_content = f.read()
                
                documents = self.load_document(file_content, file_path)
                
                if documents:
                    print(f"Loaded {len(documents)} documents from {file_path}")
                    all_documents.extend(documents)
        
        if all_documents:
            print(f"Processing {len(all_documents)} documents...")
            chunks = self.process_documents(all_documents, metadata)
            print(f"Created {len(chunks)} chunks")
            return chunks
        
        return []

    async def extract_images_from_document(self, file_content: bytes, file_name: str, file_id: str) -> List[Dict[str, Any]]:
        """문서에서 이미지 추출 (PDF, DOCX, PPTX 지원) - 재사용 가능한 버전"""
        extracted_images = []
        
        print(f"Starting image extraction for file: {file_name}")
        
        # 파일 확장자 확인
        file_extension = Path(file_name).suffix.lower()
        
        if file_extension == '.pdf':
            print("Processing PDF file for image extraction...")
            extracted_images = await self._extract_images_from_pdf(file_content, file_name, file_id)
        elif file_extension == '.docx':
            print("Processing DOCX file for image extraction...")
            extracted_images = await self._extract_images_from_docx(file_content, file_name, file_id)
        elif file_extension == '.pptx':
            print("Processing PPTX file for image extraction...")
            extracted_images = await self._extract_images_from_pptx(file_content, file_name, file_id)
        else:
            print(f"File type {file_extension} not supported for image extraction")
            
        print(f"Extracted {len(extracted_images)} images from {file_name}")
        return extracted_images

    async def _extract_images_from_pdf(self, file_content: bytes, file_name: str, file_id: str) -> List[Dict[str, Any]]:
        """PDF 파일에서 이미지 추출"""
        extracted_images = []
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            try:
                doc = fitz.open(tmp_path)
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    image_list = page.get_images()
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image['image']
                            
                            # 이미지 메타데이터
                            image_metadata = {
                                'format': base_image.get('ext', 'png'),
                                'source_path': file_name,
                                'page_number': page_num + 1,
                                'image_index': img_index
                            }
                            
                            # 이미지 이름 생성
                            image_name = f"{file_id}_page{page_num+1}_img{img_index}.{image_metadata['format']}"
                            
                            extracted_images.append({
                                'image_id': f"{file_id}_page{page_num+1}_img{img_index}",
                                'image_name': image_name,
                                'image_data': image_bytes,
                                'metadata': image_metadata
                            })
                            
                        except Exception as e:
                            print(f"Error extracting image {img_index} from page {page_num + 1}: {e}")
                            continue
                
                doc.close()
                
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            print(f"Error processing PDF file {file_name}: {e}")
            
        return extracted_images

    async def _extract_images_from_docx(self, file_content: bytes, file_name: str, file_id: str) -> List[Dict[str, Any]]:
        """DOCX 파일에서 이미지 추출"""
        extracted_images = []
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name
            
            try:
                with zipfile.ZipFile(tmp_path, 'r') as docx_zip:
                    image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                    
                    for img_index, image_path in enumerate(image_files):
                        try:
                            # 이미지 데이터 추출
                            image_data = docx_zip.read(image_path)
                            
                            # 파일 확장자 추출
                            file_extension = Path(image_path).suffix.lower()
                            if not file_extension:
                                # MIME 타입으로 판단
                                if image_data.startswith(b'\xff\xd8\xff'):
                                    file_extension = '.jpg'
                                elif image_data.startswith(b'\x89PNG\r\n\x1a\n'):
                                    file_extension = '.png'
                                elif image_data.startswith(b'GIF8'):
                                    file_extension = '.gif'
                                else:
                                    file_extension = '.png'
                            
                            # 이미지 이름 생성
                            image_name = f"{file_id}_img{img_index+1}{file_extension}"
                            
                            # 이미지 메타데이터 생성
                            image_metadata = {
                                'format': file_extension[1:],
                                'source_path': image_path,
                                'image_index': img_index + 1
                            }
                            
                            extracted_images.append({
                                'image_id': f"{file_id}_img{img_index+1}",
                                'image_name': image_name,
                                'image_data': image_data,
                                'metadata': image_metadata
                            })
                            
                        except Exception as e:
                            print(f"Error extracting image {image_path}: {e}")
                            continue
                            
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            print(f"Error processing DOCX file {file_name}: {e}")
            
        return extracted_images

    async def _extract_images_from_pptx(self, file_content: bytes, file_name: str, file_id: str) -> List[Dict[str, Any]]:
        """PPTX 파일에서 이미지 추출"""
        extracted_images = []
        
        try:
            print(f"Starting PPTX image extraction for {file_name}")
            
            # 임시 파일로 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pptx') as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name
            
            try:
                # PPTX는 ZIP 파일이므로 압축 해제
                with zipfile.ZipFile(tmp_path, 'r') as pptx_zip:
                    # 이미지 파일들 찾기
                    image_files = [f for f in pptx_zip.namelist() if f.startswith('ppt/media/')]
                    print(f"Found {len(image_files)} image files in PPTX")
                    
                    for img_index, image_path in enumerate(image_files):
                        try:
                            print(f"Processing image {img_index + 1}: {image_path}")
                            
                            # 이미지 데이터 추출
                            image_data = pptx_zip.read(image_path)
                            
                            # 파일 확장자 추출
                            file_extension = Path(image_path).suffix.lower()
                            if not file_extension:
                                # 확장자가 없는 경우 MIME 타입으로 판단
                                if image_data.startswith(b'\xff\xd8\xff'):
                                    file_extension = '.jpg'
                                elif image_data.startswith(b'\x89PNG\r\n\x1a\n'):
                                    file_extension = '.png'
                                elif image_data.startswith(b'GIF8'):
                                    file_extension = '.gif'
                                else:
                                    file_extension = '.png'  # 기본값
                            
                            print(f"Detected image format: {file_extension}")
                            
                            # 이미지 이름 생성
                            image_name = f"{file_id}_slide_img{img_index+1}{file_extension}"
                            
                            # 이미지 메타데이터 생성
                            image_metadata = {
                                'format': file_extension[1:],  # 확장자에서 점 제거
                                'source_path': image_path,
                                'slide_index': img_index + 1
                            }
                            
                            extracted_images.append({
                                'image_id': f"{file_id}_slide_img{img_index+1}",
                                'image_name': image_name,
                                'image_data': image_data,
                                'metadata': image_metadata
                            })
                            
                            print(f"Successfully extracted image {img_index + 1}")
                            
                        except Exception as e:
                            print(f"Error extracting image {image_path}: {e}")
                            continue
                            
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            print(f"Error processing PPTX file {file_name}: {e}")
            
        return extracted_images

# Example usage
if __name__ == "__main__":
    processor = DocumentProcessor()
    documents = processor.process_directory("./documents")
    print(f"Processed {len(documents)} document chunks") 